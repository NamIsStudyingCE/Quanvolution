# -*- coding: utf-8 -*-
"""build_thesis_p1.py — Luận văn FULL draft, phần 1 (v2):
front matter + Tóm tắt + Ch1 + Ch2 + Ch3.
v2 fixes theo audit Gemini: bảng 3.1/3.2/3.3 thành bảng Word thật; nhúng Hình 3.2;
tách Section để đánh số trang từ Tóm tắt; ô ký CBPB trên bìa phụ; caption nguồn."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path('GD4/KLTN_draft_full.docx')
FIG = Path('PAPER/figures')
TNR = 'Times New Roman'

doc = Document()
sec = doc.sections[0]
sec.top_margin, sec.bottom_margin = Cm(3), Cm(3.5)
sec.left_margin, sec.right_margin = Cm(3.5), Cm(2)
st = doc.styles['Normal']
st.font.name = TNR; st.font.size = Pt(13)
st.paragraph_format.line_spacing = 1.5

def para(text, size=13, bold=False, align='justify', italic=False, after=6):
    p = doc.add_paragraph()
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'justify': WD_ALIGN_PARAGRAPH.JUSTIFY, 'right': WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.name = TNR; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    return p

def h1(text): return para(text, 14, bold=True, align='left', after=10)
def h2(text): return para(text, 13, bold=True, align='left', after=6)
def caption(text): return para(text, 12, bold=True, align='center', italic=True, after=10)

def table(headers, rows, cap_text, size=11):
    caption(cap_text)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.name = TNR; r.font.size = Pt(size); r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = ''
            r = cells[j].paragraphs[0].add_run(str(v))
            r.font.name = TNR; r.font.size = Pt(size)

def pic(path, width_cm=15.5, cap=None):
    if Path(path).exists():
        doc.add_picture(str(path), width=Cm(width_cm))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        print('MISSING FIG:', path)
    if cap: caption(cap)

# ================= FRONT MATTER (Section 1 — không đánh số trang) =================
para('ĐẠI HỌC QUỐC GIA TP. HỒ CHÍ MINH', 14, bold=True, align='center')
para('TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN', 16, bold=True, align='center')
para('KHOA KỸ THUẬT MÁY TÍNH', 16, bold=True, align='center')
para('')
para('<HỌ TÊN SINH VIÊN> – <MSSV>', 14, bold=True, align='center')
para('')
para('KHÓA LUẬN TỐT NGHIỆP', 18, bold=True, align='center')
para('NGHIÊN CỨU VÀ ỨNG DỤNG LỚP TÍCH CHẬP LƯỢNG TỬ (QUANVOLUTION) TRONG PHÂN LOẠI ẢNH Y TẾ MEDMNIST, SO SÁNH CÔNG BẰNG VỚI KIẾN TRÚC CNN CỔ ĐIỂN', 16, bold=True, align='center')
para('Symmetrical Empirical Evaluation of Trainable versus Fixed Quanvolutional Filters in Medical Image Classification: A Rigorous, Reproducible Benchmark on MedMNIST', 13, italic=True, align='center')
para('')
para('KỸ SƯ NGÀNH KỸ THUẬT MÁY TÍNH', 14, bold=True, align='center')
para('TP. HỒ CHÍ MINH, 2026', 14, bold=True, align='center')
doc.add_page_break()
para('ĐẠI HỌC QUỐC GIA TP. HỒ CHÍ MINH', 14, bold=True, align='center')
para('TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN', 16, bold=True, align='center')
para('KHOA KỸ THUẬT MÁY TÍNH', 16, bold=True, align='center')
para('')
para('<HỌ TÊN SINH VIÊN> – <MSSV>', 14, bold=True, align='center')
para('KHÓA LUẬN TỐT NGHIỆP', 18, bold=True, align='center')
para('NGHIÊN CỨU VÀ ỨNG DỤNG LỚP TÍCH CHẬP LƯỢNG TỬ (QUANVOLUTION) TRONG PHÂN LOẠI ẢNH Y TẾ MEDMNIST, SO SÁNH CÔNG BẰNG VỚI KIẾN TRÚC CNN CỔ ĐIỂN', 16, bold=True, align='center')
para('GIẢNG VIÊN HƯỚNG DẪN', 14, bold=True, align='center')
para('TS. Nguyen Duy Xuan Bach', 13, align='center')
para('TP. HỒ CHÍ MINH, 2026', 14, bold=True, align='center')
para('')
para('Chữ ký xác nhận của Cán bộ phản biện (CBPB) sau bảo vệ — theo yêu cầu chỉnh sửa của Hội đồng:', 12, italic=True, align='justify')
para('………………………………………', 13, align='right')
doc.add_page_break()
para('THÔNG TIN HỘI ĐỒNG CHẤM KHÓA LUẬN TỐT NGHIỆP', 16, bold=True, align='center')
para('Hội đồng chấm khóa luận tốt nghiệp, thành lập theo Quyết định số ……… ngày ……… của Hiệu trưởng Trường Đại học Công nghệ Thông tin.', align='justify')
doc.add_page_break()
para('LỜI CẢM ƠN', 16, bold=True, align='center')
para('<Điền: Lời cảm ơn TS. Nguyen Duy Xuan Bach — GVHD, Khoa Kỹ thuật Máy tính UIT; quý thầy cô khoa; gia đình, bạn bè.>', align='justify')
doc.add_page_break()
para('MỤC LỤC', 16, bold=True, align='center')
para('<Chèn Table of Contents tự động của Word sau khi hoàn tất nội dung (References → Table of Contents).>')
doc.add_page_break()
para('DANH MỤC HÌNH', 16, bold=True, align='center')
for line in [
    'Hình 3.1: Sơ đồ pipeline Quanvolution đối xứng (mạch 4-qubit + head cổ điển)',
    'Hình 3.2: So sánh feature map — tích chập cổ điển và expectation value lượng tử',
    'Hình 3.3: Kiến trúc phần mềm hệ thống (4 tầng: dữ liệu → mô hình → thí nghiệm → đầu ra)',
    'Hình 4.1: Kết quả 10-seed BreastMNIST trên 6 metrics (5 mô hình)',
    'Hình 4.2: Kết quả 10-seed OCTMNIST trên 6 metrics (6 mô hình)',
    'Hình 4.3: Đường hội tụ train/val trên BreastMNIST và OCTMNIST',
    'Hình 4.4: Quỹ đạo góc quay θ(t) của mạch trainable_strongly qua 20 epochs',
    'Hình 4.5: Động học chuẩn gradient L2 (seed-mean ± std)',
    'Hình 4.6: Biểu đồ circuit ablation 6 cấu hình trên BreastMNIST (GĐ2)',
]:
    para(line)
doc.add_page_break()
para('DANH MỤC BẢNG', 16, bold=True, align='center')
for line in [
    'Bảng 3.1: Ba họ ansatz và số tham số tương ứng',
    'Bảng 3.2: Phân rã tham số giữa feature extractor và classifier head',
    'Bảng 3.3: Thành phần mã nguồn và vai trò trong pipeline',
    'Bảng 3.4: Cấu hình huấn luyện chi tiết (hyperparameter)',
    'Bảng 4.1: Kết quả 10-seed BreastMNIST (mean ± sample std, CI 95%)',
    'Bảng 4.2: Kết quả 10-seed OCTMNIST (mean ± sample std, CI 95%)',
    'Bảng 4.3: Kiểm định thống kê các cặp so sánh then chốt',
    'Bảng 4.4: Circuit ablation — ROC-AUC 6 cấu hình mạch',
    'Bảng 4.5: Độ trễ suy luận CPU và chi phí tính toán',
]:
    para(line)
doc.add_page_break()
para('DANH MỤC TỪ VIẾT TẮT', 16, bold=True, align='center')
for line in ['QML: Quantum Machine Learning', 'QNN: Quanvolutional Neural Network',
             'CNN: Convolutional Neural Network', 'VQC: Variational Quantum Circuit',
             'NISQ: Noisy Intermediate-Scale Quantum', 'CAD: Computer-Aided Diagnosis',
             'ROC-AUC: Area Under the Receiver Operating Characteristic Curve',
             'PR-AUC: Area Under the Precision-Recall Curve', 'MCC: Matthews Correlation Coefficient',
             'CI: Confidence Interval', 'BN: Batch Normalization']:
    para(line)

# ================= SECTION 2 — bắt đầu đánh số trang từ Tóm tắt =================
new_sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
new_sec.top_margin, new_sec.bottom_margin = Cm(3), Cm(3.5)
new_sec.left_margin, new_sec.right_margin = Cm(3.5), Cm(2)
new_sec.footer.is_linked_to_previous = False
fp = new_sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE')
_r = OxmlElement('w:r'); _t = OxmlElement('w:t'); _t.text = '1'
_r.append(_t); fld.append(_r); fp._p.append(fld)
pg = OxmlElement('w:pgNumType'); pg.set(qn('w:start'), '1')
new_sec._sectPr.append(pg)

para('TÓM TẮT KHÓA LUẬN', 16, bold=True, align='center')
para('Chẩn đoán hỗ trợ bằng máy học (CAD) trên ảnh y tế thường phải đối mặt với hai ràng buộc cùng lúc: '
     'dữ liệu có nhãn khan hiếm và mất cân bằng lớp, trong khi các kiến trúc CNN cổ điển với hàng triệu '
     'tham số dễ rơi vào overfitting trên dữ liệu nhỏ. Máy học lượng tử (QML) trong kỷ nguyên NISQ hứa hẹn '
     'một hướng tiếp cận khác: ánh xạ dữ liệu vào không gian Hilbert via mạch lượng tử tham số hóa. '
     'Tuy nhiên, phần lớn nghiên cứu hiện có thiếu tính công bằng trong so sánh: baseline cổ điển chưa '
     'được huấn luyện nghiêm túc, tầng phân loại không được cô lập, và hiếm khi có đánh giá đa seed '
     'kèm kiểm định thống kê.', align='justify')
para('Khóa luận xây dựng một khung đánh giá đối xứng 1:1 trên hai bộ dữ liệu y tế chuẩn MedMNIST v2 — '
     'BreastMNIST (780 ảnh siêu âm vú, nhị phân, lệch lớp) và OCTMNIST (subset cân bằng 5,000 ảnh võng '
     'mạc, 4 lớp) — so sánh ba nhóm mô hình: mạch tích chập lượng tử tĩnh (fixed), mạch tích chập lượng '
     'tử tự học (trainable), và baseline CNN cổ điển đối xứng tối thiểu. Toàn bộ thí nghiệm chuẩn hóa '
     'trên 10 hạt giống độc lập × 20 epochs, đánh giá 6 metrics lâm sàng, kiểm chứng bằng paired t-test, '
     'Wilcoxon signed-rank, khoảng tin cậy 95% và Cohen\u2019s d.', align='justify')
para('Kết quả chính: (1) Ưu thế lượng tử phụ thuộc chặt chẽ vào chế độ dữ liệu — trên BreastMNIST, mạch '
     'tĩnh basic_L2 đạt ROC-AUC cao nhất 0.8521 ± 0.0095 (vượt CNN 0.8336 ± 0.0259 với p = 0.0298) và '
     'mạch strongly_L2 đạt PR-AUC cao nhất 0.9182 ± 0.0071 với độ lệch chuẩn nhỏ hơn khoảng 2.7 lần; '
     'ngược lại trên OCTMNIST, CNN cổ điển áp đảo toàn bộ cấu hình lượng tử (0.7505 ± 0.0240, d = +2.108). '
     '(2) Mạch tĩnh 0 tham số cung cấp Quantum Inductive Bias mạnh mẽ trên dữ liệu nhỏ. (3) Khả năng tự '
     'học của mạch chỉ mang tính cục bộ trong cùng họ ansatz: mạch trainable_strongly vượt fixed_strongly '
     '(Δ = +0.0232, d = +1.050) nhưng chỉ hòa với một mạch ngẫu nhiên được chọn phù hợp. Chi phí: mô phỏng '
     'CPU chậm ~710 lần so với tích chập cổ điển, đòi hỏi chiến lược precompute feature maps.', align='justify')
para('Khóa luận đóng góp: một khung benchmark đối xứng, khả tái lập hoàn toàn; phân định thực nghiệm '
     'ranh giới ứng dụng của quanvolution; và bộ tài liệu kiểm định số liệu tự động.', align='justify')
para('Từ khóa: Quantum Machine Learning; Quanvolutional Neural Networks; MedMNIST; Medical Image '
     'Classification; Quantum Inductive Bias; Reproducible Benchmark', italic=True)
doc.add_page_break()

# ================= CHƯƠNG 1 =================
h1('CHƯƠNG 1. MỞ ĐẦU')
h2('1.1. Lý do chọn đề tài')
para('Chẩn đoán hỗ trợ bằng máy học (Computer-Aided Diagnosis — CAD) trên ảnh y tế — siêu âm, chụp cắt '
     'lớp quang học (OCT), tổ chức bệnh học — là một trong những hướng ứng dụng thành công nhất của học '
     'sâu hiện đại [12], [7]. Tuy nhiên, trong bối cảnh lâm sàng thực tế, dữ liệu có nhãn thường khan '
     'hiếm, quá trình gán nhãn đòi hỏi chuyên môn y khoa, và phân bố lớp thường lệch nặng (ví dụ số ca '
     'ác tính ít hơn nhiều so với lành tính). Trong khi đó, các kiến trúc CNN cổ điển với hàng triệu '
     'tham số được thiết kế cho vùng dữ liệu quy mô web (hàng triệu ảnh), khi áp đặt lên vùng dữ liệu nhỏ '
     'chúng dễ overfitting và mất khả năng tổng quát hóa [17].', align='justify')
para('Máy học lượng tử (Quantum Machine Learning — QML) trong kỷ nguyên NISQ (Noisy Intermediate-Scale '
     'Quantum) đưa ra một hướng tiếp cận khác biệt về mặt biểu diễn: mã hóa dữ liệu vào không gian '
     'trạng thái Hilbert 2^N chiều qua mạch lượng tử tham số hóa [4], [5]. Năm 2020, Henderson et al. [8] '
     'đề xuất Quanvolutional Neural Network (Quanvolution): dùng mạch lượng tử nhỏ như một bộ lọc không '
     'gian trượt trên ảnh, biến đổi từng patch cục bộ thành các kênh feature map. Giả thuyết nghiên cứu '
     'là mạch lượng tử có thể tạo ra một "thiên kiến quy nạp lượng tử" (Quantum Inductive Bias) — khai '
     'thác siêu vị và rối lượng tử để chụp các đặc trưng topo mà tích chập tuyến tính cổ điển khó tiếp cận [10].', align='justify')
para('Tuy vậy, khảo sát literature cho thấy ba khoảng trống phương pháp luận: (L1) baseline cổ điển '
     'thường yếu hoặc không đối xứng tham số, khiến không thể quy kết hiệu năng cho thành phần lượng tử; '
     '(L2) đánh giá thường chỉ trên 1–3 lần chạy, thiếu khoảng tin cậy và kiểm định ý nghĩa; (L3) tranh '
     'luận mạch tĩnh hay mạch tự học tối ưu chưa được lượng hóa trên cùng một thiết lập. Đề tài khóa '
     'luận này được chọn để trả lời các khoảng trống đó bằng một khung benchmark đối xứng, đa seed và '
     'khả tái lập hoàn toàn.', align='justify')
h2('1.2. Mục tiêu nghiên cứu')
para('Đề tài theo đuổi bốn mục tiêu: (1) xây dựng ≥ 1 mô hình quanvolution hoạt động end-to-end trên '
     'bộ dữ liệu y tế chuẩn MedMNIST; (2) thiết lập so sánh công bằng với baseline CNN cổ điển — hai '
     'phía có tầng phân loại giống hệt nhau và tổng tham số cùng cỡ; (3) đánh giá đa seed (10 hạt '
     'giống) với 6 metrics phù hợp dữ liệu y tế lệch lớp, kèm kiểm định thống kê kép; (4) phân tích '
     'trung thực ranh giới hiệu quả của quanvolution theo chế độ dữ liệu và theo vai trò mạch tĩnh/tự học.', align='justify')
h2('1.3. Đối tượng và phạm vi nghiên cứu')
para('Đối tượng: lớp tích chập lượng tử 4-qubit với mã hóa góc (angle encoding) và phép đo kỳ vọng '
     'Pauli-Z, trên hai bộ dữ liệu BreastMNIST và OCTMNIST của MedMNIST v2 [19]. Phạm vi: ảnh 28×28 '
     'grayscale; patch 2×2 stride 2; mô phỏng statevector không nhiễu (default.qubit); OCTMNIST dùng '
     'subset cân bằng 5,000 ảnh từ bộ ~97,000 ảnh gốc để đảm bảo khả năng chạy đa seed trên CPU. '
     'Phần cứng lượng tử vật lý ngoài phạm vi của khóa luận.', align='justify')
h2('1.4. Phương pháp nghiên cứu')
para('Phương pháp thực nghiệm có đối chứng: huấn luyện và đánh giá năm cấu hình (Classical CNN, Fixed '
     'Basic, Fixed Strongly, Trainable Basic, Trainable Strongly, kèm quán quân tĩnh random_L1 từ giai '
     'đoạn khảo sát) trên cùng pipeline, cùng protocol và cùng bộ hạt giống; kiểm định ý nghĩa bằng '
     'paired t-test và Wilcoxon signed-rank; đo kích thước hiệu ứng bằng Cohen\u2019s d; báo cáo '
     'mean ± sample std và khoảng tin cậy 95%.', align='justify')
h2('1.5. Đóng góp của đề tài')
para('Khóa luận có bốn đóng góp chính: (C1) khung benchmark đối xứng 1:1 với ma trận so sánh ba tầng; '
     '(C2) lượng hóa tham số và chi phí phần cứng của kernel lượng tử (0 tham số tĩnh so với 24 tham số '
     'tự học; 220 ms so với 0.31 ms mỗi ảnh trên CPU); (C3) phân định thực nghiệm ranh giới hiệu quả '
     'theo chế độ dữ liệu; (C4) kiểm tra động học gradient như một sanity check Barren Plateaus.', align='justify')
h2('1.6. Cấu trúc luận văn')
para('Luận văn gồm sáu chương. Chương 2 trình bày tổng quan lý thuyết và các công trình liên quan. '
     'Chương 3 mô tả phương pháp: kiến trúc đối xứng, các biến thể ansatz, hồ sơ thiết kế phần mềm và '
     'kiểm chứng đạo hàm. Chương 4 trình bày thiết lập thí nghiệm, kết quả trên hai dataset, động học '
     'tối ưu, ablation và chi phí tính toán. Chương 5 thảo luận và kết luận. Chương 6 nêu hướng phát triển.', align='justify')
doc.add_page_break()

# ================= CHƯƠNG 2 =================
h1('CHƯƠNG 2. TỔNG QUAN')
h2('2.1. Học sâu trong phân tích ảnh y tế')
para('Học sâu đã trở thành trụ cột của phân tích ảnh y tế hiện đại, từ phân loại tổ chức bệnh học, '
     'phát hiện tổn thương võng mạc đến phân đoạn khối u [12]. Esteva et al. [7] chỉ ra rằng các mô '
     'hình sâu có thể đạt độ chính xác ngang chuyên gia y khoa trên nhiều tác vụ da liễu và chẩn đoán '
     'hình ảnh. Tuy nhiên, thành công đó gắn liền với vùng dữ liệu lớn: ImageNet quy mô hàng triệu ảnh '
     'cho phép CNN học được phổ đặc trưng tổng quát. Trên dữ liệu y tế, ràng buộc ngược lại xuất hiện — '
     'gán nhãn tốn chi phí chuyên môn, phân bố lớp lệch tự nhiên — khiến tăng cường dữ liệu (data '
     'augmentation) [17] và các kiến trúc gọn nhẹ trở thành yêu cầu bắt buộc thay vì tùy chọn. Điều này '
     'đặt ra câu hỏi nghiên cứu: ngoài tăng dữ liệu, có tồn tại một "biến đổi đặc trưng" có cấu trúc '
     'nào đó giúp head phân loại học tốt hơn trên vùng dữ liệu nhỏ hay không — và biến đổi lượng tử là '
     'một ứng viên.', align='justify')
h2('2.2. Máy học lượng tử trong kỷ nguyên NISQ')
para('QML nghiên cứu việc dùng mạch lượng tử tham số hóa (Variational Quantum Circuit — VQC) như một '
     'khối tính toán học được trong mô hình lai lượng tử–cổ điển [4], [5]. Trong kỷ nguyên NISQ, các '
     'máy lượng tử có quy mô vài trăm qubit nhưng chưa hiệu chỉnh lỗi, nên các thuật toán khả thi phải '
     'nông, ít cổng và chịu được nhiễu. Một trục tranh luận quan trọng được Schuld và Killoran [15] nêu '
     'ra: mục tiêu "đánh bại cổ điển" (quantum advantage) có thể không phải thước đo phù hợp cho QML '
     'trong giai đoạn hiện tại; giá trị thực nằm ở việc hiểu khi nào và vì sao mô hình lượng tử mang lại '
     'lợi thế, và thước đo đánh giá phải công bằng. Huang et al. [10] chứng minh bằng lý thuyết rằng '
     'lợi thế của mô hình lượng tử phụ thuộc chặt vào bản chất dữ liệu — mệnh đề "power of data" — '
     'trực tiếp truyền cảm hứng cho thiết kế hai chế độ dữ liệu của đề tài này. Kübler et al. [11] chỉ '
     'ra rằng quantum kernel có thiên kiến quy nạp riêng ưu thế với dữ liệu có cấu trúc đối xứng.', align='justify')
h2('2.3. Quanvolutional Neural Network')
para('Henderson et al. [8] giới thiệu Quanvolution: một mạch lượng tử nhỏ đóng vai trò bộ lọc tích chập '
     'phi tuyến. Với ảnh đầu vào, một cửa sổ patch (ví dụ 2×2) được mã hóa vào các qubit bằng cổng quay '
     'góc tỉ lệ cường độ pixel; mạch biến phân biến đổi trạng thái; các kỳ vọng Pauli-Z trên từng qubit '
     'được đo để tạo thành các kênh feature map output. So với tích chập cổ điển, quanvolution có ba điểm '
     'khác biệt về bản chất: (i) biến đổi phi tuyến xuất phát tự nhiên từ phép đo lượng tử; (ii) không '
     'gian đặc trưng hàm ý 2^N chiều; (iii) tham số của mạch có thể cố định (random/fixed) hoặc tự học '
     '(trainable) — trục so sánh trung tâm của khóa luận này.', align='justify')
h2('2.4. Phân biệt với Quantum Convolutional Neural Network (QCNN)')
para('Cần phân biệt rõ quanvolution với QCNN của Cong et al. [6]. QCNN là mạng lượng tử thuần chủng: '
     'toàn bộ dữ liệu được mã hóa vào các qubit, tầng tích chập và pooling đều là các mạch lượng tử, và '
     'đầu ra cũng là phép đo lượng tử — thiết kế ban đầu cho bài toán nhận diện pha vật chất, không '
     'tương thích trực tiếp với lưới ảnh 2D thông thường. Ngược lại, quanvolution là kiến trúc lai: chỉ '
     'tầng trích xuất đặc trưng cục bộ là lượng tử, phần còn lại của mạng hoàn toàn cổ điển. Sự khác biệt '
     'này quyết định quanvolution khả thi trên hạ tầng mô phỏng CPU hiện tại với ảnh 28×28, và cũng là '
     'lý do đề tài chọn quanvolution thay vì QCNN.', align='justify')
h2('2.5. Các nghiên cứu liên quan trên dữ liệu y tế')
para('Azevedo et al. [2] ứng dụng quantum transfer learning cho phát hiện ung thư vú trên bộ mammography '
     'BCDR (825 ảnh, nhị phân). Kiến trúc dùng ResNet18 pretrained trích 512 đặc trưng, nạp vào mạch '
     'lượng tử 4-qubit dạng dressed quantum circuit, thử nghiệm cả trên simulator và máy IBM ibm_lagos. '
     'Đóng góp chính là chứng minh tính khả thi của transfer learning lượng tử trên dữ liệu vú; điểm '
     'yếu so với đề tài này là chỉ đánh giá trên một split duy nhất, không đa seed, và baseline so sánh '
     'không được đối xứng tham số với mạch lượng tử.', align='justify')
para('Matondo-Mvula và Elleithy [13] là công trình gần đề tài nhất: quanvolution 9-qubit với kernel '
     '3×3 (15 tham số học mỗi kernel) trên chính BreastMNIST, hạ độ phân giải 28×28 về 14×14 để giảm '
     'chi phí. Kết quả QCNN 67% test accuracy, thấp hơn CNN baseline 83.33% validation. Công trình này '
     'xác nhận hai điều: quanvolution chạy được trên dữ liệu siêu âm vú thực tế, và việc thiếu đa seed '
     'cùng baseline đối xứng khiến kết luận không thể suy rộng — chính là hai giới hạn đề tài này khắc phục.', align='justify')
para('Vu, Le và Pham [18] khảo sát hệ thống ảnh hưởng của các thành phần feature trong quanvolution '
     'cho phân loại ảnh tổng quát, chỉ ra rằng lựa chọn số qubit và cấu trúc mạch ảnh hưởng đáng kể '
     'đến chất lượng đặc trưng. Công trình đặt nền phương pháp luận cho việc ablation cấu hình mạch — '
     'đề tài kế thừa tư tưởng này ở quy mô GĐ2 (6 cấu hình) và mở rộng sang cả chế độ dữ liệu.', align='justify')
para('Hoang et al. [9] đề xuất lớp truyền đặc trưng (feature propagation layer) cho HQCNN đa lớp, '
     'giải quyết vấn đề batching khi tích hợp mạch lượng tử vào CNN. Altares-López et al. [1] tự động '
     'thiết kế quantum feature maps bằng thuật toán di truyền, cho thấy không gian thiết kế mạch rất '
     'lớn và việc chọn mạch ngẫu nhiên/cố định có cấu trúc là một quyết định thiết kế quan trọng — '
     'cơ sở lý thuyết cho trục so sánh fixed-vs-trainable của đề tài.', align='justify')
para('Kübler et al. [11] chứng minh quantum kernel có thiên kiến quy nạp phụ thuộc cấu trúc dữ liệu; '
     'Schuld và Killoran [15] đưa ra góc nhìn phản biện rằng quantum advantage không phải mục tiêu '
     'duy nhất đáng theo đuổi. Hai công trình này định hình triết lý đánh giá của đề tài: thay vì '
     'tìm kiếm lợi thế tuyệt đối, hãy xác định chính xác điều kiện dữ liệu mà thành phần lượng tử '
     'đóng góp tích cực — và đo bằng protocol công bằng nhất có thể.', align='justify')
h2('2.6. Vị trí của đề tài')
para('Đề tài khác biệt ở bốn điểm: (i) baseline CNN được thiết kế đối xứng 1:1 — cùng head, cùng số '
     'chiều đặc trưng 784, tổng tham số chênh lệch không quá 24; (ii) 10 seeds độc lập với kiểm định '
     'kép và kích thước hiệu ứng; (iii) hai chế độ dữ liệu đối lập (nhỏ/lệch lớp và lớn/đa lớp) để phân '
     'định ranh giới hiệu quả; (iv) kết quả trung thực theo tinh thần "không cần quantum thắng, cần biết '
     'khi nào và vì sao" — nhất quán với lập luận của Schuld và Killoran [15].', align='justify')
doc.add_page_break()

# ================= CHƯƠNG 3 =================
h1('CHƯƠNG 3. PHƯƠNG PHÁP NGHIÊN CỨU')
h2('3.1. Kiến trúc pipeline đối xứng')
para('Pipeline end-to-end gồm bốn giai đoạn tuần tự: (1) phân mảnh — ảnh 28×28×1 được chia thành 196 '
     'patch 2×2 không chồng lấn (stride 2); (2) kernel lượng tử 4-qubit — mỗi patch được mã hóa góc '
     'R_Y(π·x_i) vào 4 qubit, qua mạch biến phân U(θ), đo kỳ vọng Pauli-Z trên từng qubit; (3) feature '
     'map lượng tử — đầu ra tổ chức thành tensor 4×14×14 (784 chiều); (4) head phân loại đối xứng — '
     'BatchNorm2d(4) → ReLU → Linear(784, K) với K = 2 (BreastMNIST) hoặc K = 4 (OCTMNIST). Baseline '
     'CNN cổ điển thay giai đoạn (2) bằng một lớp Conv2D(1→4, kernel 2×2, stride 2, có bias) và giữ '
     'nguyên toàn bộ giai đoạn (4).', align='justify')
pic(FIG / 'Fig1_quanvolution_pipeline.png', 15.5,
    'Hình 3.1: Sơ đồ pipeline Quanvolution đối xứng — head cổ điển giống hệt nhau giữa hai phía '
    '(Nguồn: tác giả tổng hợp từ tài liệu dự án)')
h2('3.2. Công thức mã hóa và phép đo')
para('Patch x = (x₀, x₁, x₂, x₃), x_i ∈ [0,1] được mã hóa: |ψ(x)⟩ = ⊗ᵢ R_Y(π·x_i)|0⟩. Mạch biến phân '
     'U(θ) biến đổi |Φ(x,θ)⟩ = U(θ)|ψ(x)⟩. Feature map ở tọa độ (u,v), kênh i: F_i(u,v) = '
     '⟨Φ(x,θ)|Z_i|Φ(x,θ)⟩ ∈ [−1,1]. Với ảnh 28×28, đầu ra là 4×14×14 = 784 chiều — đúng bằng số chiều '
     'đầu ra của baseline conv đối xứng, đảm bảo head hai phía cùng kiến trúc. Hình 3.2 so sánh trực '
     'quan feature map tạo bởi tích chập cổ điển và bởi phép đo kỳ vọng lượng tử trên cùng một ảnh.', align='justify')
pic(FIG / 'Fig2_feature_comparison.png', 15.5,
    'Hình 3.2: So sánh feature map — tích chập cổ điển và expectation value lượng tử '
    '(Nguồn: tác giả tổng hợp từ kết quả thực nghiệm đề tài)')
h2('3.3. Ba họ ansatz và đối xứng tham số')
para('Ba họ mạch được khảo sát, tóm tắt tại Bảng 3.1:', align='justify')
table(
    ['Họ ansatz', 'Cổng quay', 'Lớp rối (entangling)', 'Tham số/tầng', 'Tham số (L=2 / L=1)'],
    [
        ['Basic Entangling', 'R_Y(θ_i) — 1 trục', 'CNOT vòng q0→q1→q2→q3→q0', '4L', '8 / 4'],
        ['Strongly Entangling', 'U₃ = R_Z(ω)R_Y(θ)R_Z(φ) — 3 trục', 'Rối tuần hoàn các cặp qubit', '12L', '24 / 12'],
        ['Random (đóng băng)', 'Cổng quay 3 trục ngẫu nhiên Haar', 'Cặp CNOT ngẫu nhiên', '0 (frozen)', '0 / 0'],
    ],
    'Bảng 3.1: Ba họ ansatz và số tham số tương ứng (N_q = 4)')
para('Bảng 3.2 phân rã tham số theo thành phần. Điểm mấu chốt của thiết kế đối xứng: tổng tham số giữa '
     'baseline CNN (1,598 với K=2) và quanvolution tĩnh (1,578) chênh lệch đúng 20 tham số của lớp conv '
     'cổ điển — trong khi head phân loại (1,570 + 8 BatchNorm) giống hệt nhau. Như vậy bất kỳ chênh lệch '
     'hiệu năng nào đều có thể quy về bản chất của phép biến đổi đặc trưng, không phải dung lượng tham số.', align='justify')
table(
    ['Mô hình', 'Kernel (L=2 / L=1)', 'BatchNorm2d(4)', 'Head (K=2)', 'Head (K=4)', 'Tổng (K=2, L=2 / K=4, L=1)'],
    [
        ['Classical CNN', '20 (conv có bias)', '8', '1,570', '3,140', '1,598 / 3,168'],
        ['Fixed (mọi ansatz)', '0', '8', '1,570', '3,140', '1,578 / 3,148'],
        ['Trainable Basic', '8 / 4', '8', '1,570', '3,140', '1,586 / 3,152'],
        ['Trainable Strongly', '24 / 12', '8', '1,570', '3,140', '1,602 / 3,160'],
    ],
    'Bảng 3.2: Phân rã tham số giữa feature extractor và classifier head (Nguồn: đếm tự động từ '
    'instantiation PyTorch — measure_params_cost.py)')
h2('3.4. Hồ sơ thiết kế phần mềm')
para('Pipeline phần mềm tổ chức theo bốn mô-đun, mỗi mô-đun một trách nhiệm đơn nhất (Bảng 3.3), toàn '
     'bộ mã nguồn đặt trên GitHub với seed cố định và gắn tag phiên bản nộp hội nghị '
     '(soict-submission-v4). Luồng dữ liệu chuẩn: run_gd3.py → thí nghiệm 50 runs (BreastMNIST) / '
     '60 runs (OCTMNIST) → full_trainable_*.json → reconcile_verify.py (kiểm định số liệu) → '
     'regenerate_figs_bigfont.py (biểu đồ 300 DPI). Mọi thay đổi số liệu phải đi qua kịch bản kiểm định '
     'tự động để loại trừ lỗi sao chép tay. Hình 3.3 trực quan hóa kiến trúc bốn tầng của hệ thống.', align='justify')
pic(str(Path('GD4') / 'fig_software_architecture.png'), 15.5,
    'Hình 3.3: Kiến trúc phần mềm hệ thống — bốn tầng từ dữ liệu đến tài liệu/demo '
    '(Nguồn: tác giả thiết kế từ cấu trúc mã nguồn src/)')
table(
    ['Mô-đun', 'Thành phần chính', 'Vai trò trong pipeline'],
    [
        ['src/data', 'medmnist_loader.py, precompute_features.py', 'Nạp MedMNIST, chuẩn hóa [0,1], chia split cố định, precompute quantum features cho mạch tĩnh'],
        ['src/models', 'circuits.py, trainable_quanv.py, quantum_model.py, classical_cnn.py', 'Định nghĩa 6 mạch (3 ansatz × L=1,2), TorchLayer differentiable, hai mạng có head giống hệt nhau'],
        ['src/experiments', 'run_gd3.py, trainable_experiment.py', 'Vòng lặp 10 seeds, huấn luyện, đánh giá 6 metrics, kiểm định thống kê, xuất JSON có cấu trúc'],
        ['src/utils + src/visual', 'metrics.py, plot_gd3_dynamics.py', 'Bộ 6 metrics y tế, seed utils, vẽ biểu đồ 300 DPI'],
    ],
    'Bảng 3.3: Thành phần mã nguồn và vai trò trong pipeline (Nguồn: tài liệu dự án)')
h2('3.5. Huấn luyện và hai chiến lược tính toán')
para('Protocol huấn luyện chung: CrossEntropy loss, Adam với learning rate kép (0.001 cho tham số cổ '
     'điển, 0.01 cho góc lượng tử), batch 32, 20 epochs, chọn checkpoint theo best val ROC-AUC. Với mạch '
     'tĩnh, feature map được precompute một lần cho toàn bộ dataset (10 seeds chỉ tốn ~18 giây huấn '
     'luận head); với mạch tự học, mạch lượng tử được vi phân bằng backprop statevector tích hợp '
     'PyTorch (Bergholm et al. [3]) và các góc θ cập nhật đồng thời head với learning rate riêng.', align='justify')
caption('Bảng 3.4: Cấu hình huấn luyện chi tiết (hyperparameter)')
table(
    ['Hyperparameter', 'Giá trị', 'Ghi chú'],
    [
        ['Optimizer', 'Adam', 'Tách nhóm tham số lượng tử / cổ điển'],
        ['Learning rate (cổ điển)', '0.001', 'Head + BatchNorm'],
        ['Learning rate (lượng tử θ)', '0.01', 'Learning rate kép, góc quay hội tụ nhanh hơn'],
        ['Loss', 'CrossEntropy', 'Nhãn nhị phân / 4 lớp'],
        ['Batch size', '32', 'BreastMNIST 546 train → 18 batch/epoch'],
        ['Epochs', '20', 'Plateau hội tụ 12–15 (xem Chương 4)'],
        ['Seeds', '10 (S = {0, 42, 100, 2023, 777, 999, 1234, 5678, 1111, 2222})', 'Độc lập toàn phần (data split, init, shuffle)'],
        ['Checkpoint', 'Best val ROC-AUC', 'Đánh giá test trên checkpoint tốt nhất'],
        ['Mô phỏng', 'default.qubit (statevector, không nhiễu)', 'Analytic backprop + kiểm chứng Parameter-Shift'],
    ],
    'Bảng 3.4: Cấu hình huấn luyện chi tiết (Nguồn: tài liệu dự án — protocol đồng nhất mọi mô hình)')
para('Vòng lặp huấn luyện một epoch của mạch tự học tóm tắt bằng mã giả như sau (protocol giống hệt '
     'baseline CNN, chỉ khác khối trích xuất đặc trưng là mạch lượng tử differentiable):', align='justify')
for line in [
    'set_seed(SEED); head ← QuanvolutionClassifier(784→K); θ ← init_uniform(0, 2π)',
    'for epoch in 1..20:',
    '    for (x_img, y) in train_loader:',
    '        f ← QuanvLayer(x_img)                  # 196 patches × mạch 4-qubit (differentiable)',
    '        logits ← BN → ReLU → Linear(f)          # head cổ điển đối xứng',
    '        loss ← CrossEntropy(logits, y); backward(); Adam.step()   # lr kép 0.01(θ) / 0.001(head)',
    '    va_auc ← evaluate(val_loader)',
    '    if va_auc > best: best ← va_auc; checkpoint ← state_dict()',
    'test_metrics ← evaluate(test_loader, checkpoint=best)',
]:
    p = para(line, 11, after=1)
    p.runs[0].font.name = 'Courier New'
h2('3.6. Kiểm chứng đạo hàm')
para('Để bảo đảm tính đúng đắn vật lý của gradient analytic statevector, đạo hàm được đối chứng với '
     'Parameter-Shift Rule [16]: ∂F_i/∂θ_j = [F_i(θ_j + π/2) − F_i(θ_j − π/2)]/2. Sai lệch trung bình '
     'tuyệt đối đo được |Δ| < 4.1×10⁻⁸ — cấp độ sai số số học float, xác nhận hai đường tính gradient '
     'tương đương về mặt vật lý.', align='justify')
doc.save(str(OUT))
print('PART 1 (v2) saved:', OUT)
