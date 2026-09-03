# -*- coding: utf-8 -*-
"""build_thesis_skeleton.py — C1/C9: dựng khung luận văn UIT + đề cương chi tiết đã điền.
Format theo Phụ lục 2 (hình thức trình bày) & Phụ lục 3 (mẫu báo cáo) của UIT."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

OUT_DIR = Path('GD4')
OUT_DIR.mkdir(exist_ok=True)

TNR = 'Times New Roman'

def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin, sec.bottom_margin = Cm(3), Cm(3.5)
    sec.left_margin, sec.right_margin = Cm(3.5), Cm(2)
    st = doc.styles['Normal']
    st.font.name = TNR
    st.font.size = Pt(13)
    st.paragraph_format.line_spacing = 1.5
    return doc

def para(doc, text, size=13, bold=False, align='left', space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'justify': WD_ALIGN_PARAGRAPH.JUSTIFY, 'right': WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = TNR; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    return p

def h1(doc, text):  # Chương: bold 14
    p = para(doc, text, size=14, bold=True, align='left', space_after=10)
    return p

def h2(doc, text):  # mục: bold 13
    return para(doc, text, size=13, bold=True, space_after=6)

def body(doc, text):
    return para(doc, text, size=13, align='justify')

def page_break(doc):
    doc.add_page_break()

# ============ FILE 1: ĐỀ CƯƠNG CHI TIẾT (đã điền) ============
doc = setup_doc()
para(doc, 'ĐẠI HỌC QUỐC GIA TP. HỒ CHÍ MINH', 14, bold=True, align='center')
para(doc, 'TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN', 16, bold=True, align='center')
para(doc, 'KHOA KỸ THUẬT MÁY TÍNH', 16, bold=True, align='center')
para(doc, 'CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM', 14, bold=True, align='center')
para(doc, 'Độc Lập – Tự Do – Hạnh Phúc', 14, bold=True, align='center')
para(doc, '')
para(doc, 'ĐỀ CƯƠNG CHI TIẾT', 16, bold=True, align='center')
para(doc, '')
para(doc, 'TÊN ĐỀ TÀI TIẾNG VIỆT:  Nghiên cứu và ứng dụng lớp tích chập lượng tử (Quanvolution) '
          'trong bài toán phân loại ảnh y tế MedMNIST, so sánh công bằng với kiến trúc CNN cổ điển', 13, align='justify')
para(doc, "TÊN ĐỀ TÀI TIẾNG ANH:  Symmetrical Empirical Evaluation of Trainable versus Fixed "
          "Quanvolutional Filters in Medical Image Classification: A Rigorous, Reproducible Benchmark on MedMNIST", 13, align='justify')
para(doc, 'Cán bộ hướng dẫn:  TS. Nguyen Duy Xuan Bach — Khoa Kỹ thuật Máy tính, '
          'Trường Đại học Công nghệ Thông tin, ĐHQG-HCM', 13, align='justify')
para(doc, 'Thời gian thực hiện:  Từ ngày 11/08/2026 đến ngày 09/11/2026', 13)
para(doc, 'Sinh viên thực hiện:  <Họ tên sinh viên – MSSV>', 13)
para(doc, '')
h2(doc, 'Nội dung đề tài')
body(doc, '1. Tổng quan đề tài: Quanvolutional Neural Network (Henderson et al., 2020) dùng mạch lượng tử '
          '4-qubit làm bộ lọc 2×2 trượt trên ảnh để tạo feature map phi tuyến trong không gian Hilbert. '
          'Các nghiên cứu QML trên ảnh y tế hiện có còn ba giới hạn chính: (i) baseline cổ điển chưa được '
          'tune nghiêm túc / không đối xứng tham số với mô hình lượng tử; (ii) thiếu đánh giá đa seed và '
          'kiểm định thống kê; (iii) chưa phân định rõ vai trò của mạch tĩnh so với mạch tự học. Từ thực '
          'trạng đó, đề tài xây dựng một khung benchmark đối xứng 1:1, khả tái lập để đo lường chính xác '
          'giá trị của lớp tích chập lượng tử.')
body(doc, '2. Mục tiêu của đề tài: (i) xây dựng mô hình quanvolution end-to-end trên ≥ 1 bộ dữ liệu '
          'MedMNIST với baseline CNN cổ điển được huấn luyện nghiêm túc, đối xứng từng tham số; '
          '(ii) đánh giá đa seed (10 seeds) với 6 metrics y tế (Accuracy, Balanced Accuracy, F1, MCC, '
          'ROC-AUC, PR-AUC) kèm kiểm định kép paired t-test và Wilcoxon, CI 95%, Cohen\u2019s d; '
          '(iii) cải tiến: ma trận so sánh 3 tầng (mạch tĩnh — mạch tự học — quán quân tĩnh), phân định '
          'ranh giới hiệu quả theo chế độ dữ liệu (data-regime dependency); (iv) demo end-to-end và '
          'bản thảo bài báo quốc tế.')
body(doc, '3. Phương pháp thực hiện: mô phỏng statevector (PennyLane default.qubit); mạch tĩnh dùng '
          'chiến lược precompute feature maps; mạch tự học dùng vi phân backprop tích hợp PyTorch; '
          'head phân loại cổ điển đối xứng (BatchNorm2d → ReLU → Linear 784→K); CrossEntropy + Adam '
          '(lr kép 0.001 cổ điển / 0.01 lượng tử); kiểm tra đạo hàm đối chứng Parameter-Shift Rule.')
body(doc, '4. Các nội dung chính và giới hạn của đề tài: hai bộ dữ liệu BreastMNIST (780 ảnh, nhị phân, '
          'lệch lớp) và OCTMNIST (subset cân bằng 5,000 ảnh, 4 lớp); ảnh 28×28, patch 2×2, 4 qubits; '
          'cấu hình mạch: Basic/Strongly/Random × L=1,2. Giới hạn: mô phỏng không nhiễu (chưa chạy NISQ '
          'hardware thật), OCTMNIST dùng subset thay vì ~97,000 ảnh đầy đủ. Hệ thống demo: notebook '
          'chạy live — nạp ảnh siêu âm → trích xuất feature map bằng mạch lượng tử → dự đoán kèm xác suất.')
para(doc, '')
h2(doc, 'Kế hoạch thực hiện')
for line in [
    'GĐ0 (11/08 – 24/08) — Nền tảng & môi trường: lý thuyết quanvolution, cài đặt PennyLane/PyTorch, demo MNIST. Trạng thái: hoàn thành (mốc M1).',
    'GĐ1 (25/08 – 07/09) — Pipeline dữ liệu MedMNIST & baseline CNN công bằng, 10 seeds. Trạng thái: hoàn thành (mốc M2).',
    'GĐ2 (08/09 – 28/09) — Khảo sát mạch lượng tử tĩnh, circuit ablation 6 cấu hình. Trạng thái: hoàn thành (mốc M3).',
    'GĐ3 (29/09 – 19/10) — Ma trận 3 tầng (fixed/trainable/champion), mạch tự học, kiểm định thống kê; bản thảo bài báo nộp SOICT 2026. Trạng thái: hoàn thành (mốc M4).',
    'GĐ4 (20/10 – 09/11) — Viết luận văn, demo notebook, slide bảo vệ. Trạng thái: đang thực hiện.',
]:
    para(doc, '• ' + line, 13, align='justify')
para(doc, '')
para(doc, 'TP. HCM, ngày … tháng … năm 2026', 13, align='right')
para(doc, 'Xác nhận của CBHD (ký, ghi rõ họ tên)          Sinh viên (ký, ghi rõ họ tên)', 13)
OUT2 = OUT_DIR / 'de_cuong_chi_tiet_filled_v2.docx'
try:
    doc.save(OUT_DIR / 'de_cuong_chi_tiet_filled.docx')
except PermissionError:
    doc.save(OUT2)
    print('saved', OUT2, '(bản cũ đang mở trong Word)')
print('saved', OUT_DIR / 'de_cuong_chi_tiet_filled.docx')

# ============ FILE 2: KHUNG LUẬN VĂN ============
doc = setup_doc()

# Bìa chính (theo Phụ lục 3)
para(doc, 'ĐẠI HỌC QUỐC GIA TP. HỒ CHÍ MINH', 14, bold=True, align='center')
para(doc, 'TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN', 16, bold=True, align='center')
para(doc, 'KHOA KỸ THUẬT MÁY TÍNH', 16, bold=True, align='center')
para(doc, '')
para(doc, '<HỌ TÊN SINH VIÊN> – <MSSV>', 14, bold=True, align='center')
para(doc, '')
para(doc, 'KHÓA LUẬN TỐT NGHIỆP', 18, bold=True, align='center')
para(doc, '')
para(doc, 'NGHIÊN CỨU VÀ ỨNG DỤNG LỚP TÍCH CHẬP LƯỢNG TỬ (QUANVOLUTION) TRONG PHÂN LOẠI ẢNH Y TẾ MEDMNIST, SO SÁNH CÔNG BẰNG VỚI KIẾN TRÚC CNN CỔ ĐIỂN', 16, bold=True, align='center')
para(doc, 'Symmetrical Empirical Evaluation of Trainable versus Fixed Quanvolutional Filters in Medical Image Classification: A Rigorous, Reproducible Benchmark on MedMNIST', 13, italic=True, align='center')
para(doc, '')
para(doc, 'KỸ SƯ NGÀNH KỸ THUẬT MÁY TÍNH', 14, bold=True, align='center')
para(doc, '')
para(doc, 'TP. HỒ CHÍ MINH, 2026', 14, bold=True, align='center')
page_break(doc)

# Bìa phụ
para(doc, 'ĐẠI HỌC QUỐC GIA TP. HỒ CHÍ MINH', 14, bold=True, align='center')
para(doc, 'TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN', 16, bold=True, align='center')
para(doc, 'KHOA KỸ THUẬT MÁY TÍNH', 16, bold=True, align='center')
para(doc, '')
para(doc, '<HỌ TÊN SINH VIÊN> – <MSSV>', 14, bold=True, align='center')
para(doc, 'KHÓA LUẬN TỐT NGHIỆP', 18, bold=True, align='center')
para(doc, 'NGHIÊN CỨU VÀ ỨNG DỤNG LỚP TÍCH CHẬP LƯỢNG TỬ (QUANVOLUTION) TRONG PHÂN LOẠI ẢNH Y TẾ MEDMNIST, SO SÁNH CÔNG BẰNG VỚI KIẾN TRÚC CNN CỔ ĐIỂN', 16, bold=True, align='center')
para(doc, 'GIẢNG VIÊN HƯỚNG DẪN', 14, bold=True, align='center')
para(doc, 'TS. Nguyen Duy Xuan Bach', 13, align='center')
para(doc, 'TP. HỒ CHÍ MINH, 2026', 14, bold=True, align='center')
page_break(doc)

# Trang hội đồng
para(doc, 'THÔNG TIN HỘI ĐỒNG CHẤM KHÓA LUẬN TỐT NGHIỆP', 16, bold=True, align='center')
para(doc, 'Hội đồng chấm khóa luận tốt nghiệp, thành lập theo Quyết định số ……… ngày ……… của Hiệu trưởng Trường Đại học Công nghệ Thông tin.', 13, align='justify')
page_break(doc)

# Lời cảm ơn
para(doc, 'LỜI CẢM ƠN', 16, bold=True, align='center')
body(doc, '<Điền nội dung lời cảm ơn: GVHD TS. Nguyen Duy Xuan Bach, Khoa Kỹ thuật Máy tính UIT, gia đình, bạn bè.>')
page_break(doc)

# Mục lục (placeholder — Word tự cập nhật bằng TOC field khi chèn)
para(doc, 'MỤC LỤC', 16, bold=True, align='center')
body(doc, '<Chèn Mục lục tự động: References → Table of Contents, sau khi xong nội dung.>')
page_break(doc)

# Danh mục hình — pre-fill từ tài sản paper (đánh số lại theo Chương.số)
para(doc, 'DANH MỤC HÌNH', 16, bold=True, align='center')
for line in [
    'Hình 3.1: Sơ đồ pipeline Quanvolution đối xứng (mạch 4-qubit + head cổ điển) ',
    'Hình 3.2: Kiến trúc class/pipeline mã nguồn (circuits, trainable_quanv, quantum_model)',
    'Hình 4.1: So sánh feature map — tích chập cổ điển vs expectation value lượng tử',
    'Hình 4.2: Kết quả 10-seed BreastMNIST (6 metrics, 5 mô hình)',
    'Hình 4.3: Kết quả 10-seed OCTMNIST (6 metrics, 6 mô hình)',
    'Hình 4.4: Đường hội tụ train/val BreastMNIST và OCTMNIST',
    'Hình 4.5: Quỹ đạo góc quay θ(t) qua 20 epochs',
    'Hình 4.6: Động học chuẩn gradient L2 (chứng cứ không Barren Plateau)',
    'Hình 4.7: Biểu đồ circuit ablation 6 cấu hình (GĐ2)',
]:
    para(doc, line, 13)
page_break(doc)

# Danh mục bảng
para(doc, 'DANH MỤC BẢNG', 16, bold=True, align='center')
for line in [
    'Bảng 3.1: Phân rã tham số giữa feature extractor và classifier head',
    'Bảng 4.1: Kết quả 10-seed BreastMNIST (6 metrics ± std, CI 95%)',
    'Bảng 4.2: Kết quả 10-seed OCTMNIST (6 metrics ± std, CI 95%)',
    'Bảng 4.3: Kiểm định thống kê cặp then chốt (t-test, Wilcoxon, Cohen\u2019s d)',
    'Bảng 4.4: Circuit ablation 6 cấu hình (GĐ2)',
    'Bảng 4.5: Độ trễ suy luận CPU và chi phí tính toán',
]:
    para(doc, line, 13)
page_break(doc)

# Danh mục từ viết tắt
para(doc, 'DANH MỤC TỪ VIẾT TẮT', 16, bold=True, align='center')
for line in ['QML: Quantum Machine Learning', 'QNN: Quanvolutional Neural Network',
             'CNN: Convolutional Neural Network', 'VQC: Variational Quantum Circuit',
             'NISQ: Noisy Intermediate-Scale Quantum', 'CAD: Computer-Aided Diagnosis',
             'ROC-AUC: Area Under the Receiver Operating Characteristic Curve',
             'PR-AUC: Area Under the Precision-Recall Curve', 'MCC: Matthews Correlation Coefficient',
             'CI: Confidence Interval']:
    para(doc, line, 13)
page_break(doc)

# Tóm tắt (1-2 trang) — bắt đầu đánh số trang từ đây
para(doc, 'TÓM TẮT KHÓA LUẬN', 16, bold=True, align='center')
body(doc, '<DRAFT 1.5 trang — khung từ abstract paper mở rộng: (1) bối cảnh CAD + giới hạn QML; '
          '(2) phương pháp benchmark đối xứng 1:1, 10 seeds, 6 metrics, kiểm định kép; '
          '(3) kết quả BreastMNIST (quantum thắng ROC-AUC 0.8521±0.0095, PR-AUC 0.9182±0.0071, '
          'std nhỏ hơn ~2.7×) và OCTMNIST (CNN thắng 0.7505±0.0240, d=+2.108); '
          '(4) ba kết luận: data-regime dependency, 0-param inductive bias, trainability cục bộ; '
          '(5) từ khóa.>')
page_break(doc)

# Chương 1 MỞ ĐẦU
h1(doc, 'CHƯƠNG 1. MỞ ĐẦU')
h2(doc, '1.1. Lý do chọn đề tài')
body(doc, '<Từ Introduction paper + bối cảnh CAD, dữ liệu y tế nhỏ/lệch lớp, chi phí phần cứng lượng tử.>')
h2(doc, '1.2. Mục tiêu nghiên cứu')
body(doc, '<4 mục tiêu như đề cương chi tiết.>')
h2(doc, '1.3. Đối tượng và phạm vi nghiên cứu')
body(doc, '<Đối tượng: quanvolution 4-qubit trên MedMNIST. Phạm vi: 2 datasets, 28×28, mô phỏng không nhiễu.>')
h2(doc, '1.4. Phương pháp nghiên cứu')
body(doc, '<Benchmark đối xứng, đa seed, kiểm định thống kê kép.>')
h2(doc, '1.5. Đóng góp của đề tài')
body(doc, '<C1 benchmark đối xứng 3 tầng; C2 tham số & chi phí; C3 phân định data-regime; C4 gradient sanity check.>')
h2(doc, '1.6. Cấu trúc luận văn')
body(doc, '<Mô tả 6 chương.>')
page_break(doc)

# Chương 2 TỔNG QUAN
h1(doc, 'CHƯƠNG 2. TỔNG QUAN')
h2(doc, '2.1. Chẩn đoán hỗ trợ bằng máy học trên ảnh y tế')
h2(doc, '2.2. Máy học lượng tử trong kỷ nguyên NISQ')
h2(doc, '2.3. Quanvolutional Neural Network')
h2(doc, '2.4. Phân biệt với Quantum Convolutional Neural Network (QCNN)')
h2(doc, '2.5. Tổng quan các nghiên cứu liên quan và khoảng trống nghiên cứu')
body(doc, '<Bảng so sánh 7 công trình (đã verify) + 3 limitation L1–L3 + định vị đề tài. '
          'Nguồn: Related Work paper + docs/theory_summary.md — mở rộng mỗi mục 2–3 trang.>')

# Chương 3 PHƯƠNG PHÁP
h1(doc, 'CHƯƠNG 3. PHƯƠNG PHÁP NGHIÊN CỨU')
h2(doc, '3.1. Kiến trúc pipeline đối xứng')
body(doc, '<Paper §3.1 + Hình 3.1.>')
h2(doc, '3.2. Mạch lượng tử và các biến thể ansatz')
body(doc, '<Paper §2.1 + §3.2: angle embedding RY(πx), 3 họ mạch, công thức 1–3.>')
h2(doc, '3.3. Baseline CNN cổ điển đối xứng')
body(doc, '<Paper §3.3 + đối chiếu tham số 28 vs 8, head giống hệt.>')
h2(doc, '3.4. Hồ sơ thiết kế phần mềm')
body(doc, '<BẮT BUỘC theo quy định (sản phẩm phần mềm): sơ đồ class src/models/, luồng dữ liệu '
          'run_gd3.py, chiến lược precompute, dual learning rate. Thêm sơ đồ UML component/sequence.>')
h2(doc, '3.5. Kiểm chứng đạo hàm và tính đúng đắn')
body(doc, '<Parameter-shift vs backprop, |Δ| < 4.1×10⁻⁸ — Paper §3.4.>')

# Chương 4 KẾT QUẢ
h1(doc, 'CHƯƠNG 4. KẾT QUẢ THỰC NGHIỆM VÀ ĐÁNH GIÁ')
h2(doc, '4.1. Thiết lập thí nghiệm')
body(doc, '<Datasets, 10 seeds, 20 epochs, 6 metrics, kiểm định; môi trường PennyLane/PyTorch.>')
h2(doc, '4.2. Kết quả trên BreastMNIST')
body(doc, '<Bảng 4.1 + Hình 4.2 + diễn giải: quantum thắng ROC-AUC/PR-AUC, std nhỏ hơn ~2.7×.>')
h2(doc, '4.3. Kết quả trên OCTMNIST')
body(doc, '<Bảng 4.2 + Hình 4.3: CNN thắng áp đảo 0.7505±0.0240, d=+2.108.>')
h2(doc, '4.4. Động học tối ưu và kiểm chứng gradient')
body(doc, '<Hình 4.4–4.6: hội tụ 12–15 epochs, chuẩn gradient 0.2–0.5 seed-mean.>')
h2(doc, '4.5. Circuit ablation')
body(doc, '<Bảng 4.4 + Hình 4.7 từ GĐ2 — nội dung paper không có, mở rộng riêng cho luận văn.>')
h2(doc, '4.6. Chi phí tính toán')
body(doc, '<Bảng 4.5: 0.31 ms vs 220.22 ms, 710×; precompute 18s/10 seeds.>')

# Chương 5 KẾT LUẬN
h1(doc, 'CHƯƠNG 5. KẾT LUẬN')
h2(doc, '5.1. Kết luận')
body(doc, '<3 kết luận: data-regime dependency; 0-param inductive bias; trainability cục bộ. '
          'Ngắn gọn, không bàn luận thêm (quy định Phụ lục 2).>')
h2(doc, '5.2. Những đóng góp mới')
body(doc, '<Khung benchmark đối xứng + phân định ranh giới + bộ tài liệu tái lập.>')

# Chương 6 HƯỚNG PHÁT TRIỂN
h1(doc, 'CHƯƠNG 6. HƯỚNG PHÁT TRIỂN')
body(doc, '<GPU/cuQuantum, NISQ noise, full-scale OCTMNIST, resolution cao hơn, hybrid gating so sánh.>')

# TLTK
page_break(doc)
para(doc, 'TÀI LIỆU THAM KHẢO', 16, bold=True, align='center')
h2(doc, 'Tiếng Việt')
body(doc, '<Nếu có — để trống nếu không trích dẫn tài liệu tiếng Việt.>')
h2(doc, 'Tiếng Anh')
body(doc, '<19 references IEEE của paper (đã verify 100% — tái dùng nguyên văn, tách alphabet). '
          'Lưu ý: chỉ liệt kê tài liệu được TRÍCH DẪN trong luận văn.>')

# Phụ lục
page_break(doc)
para(doc, 'PHỤ LỤC', 16, bold=True, align='center')
h2(doc, 'Phụ lục A. Hướng dẫn cài đặt và tái hiện thực nghiệm')
body(doc, '<Từ README: pip install -r requirements.txt; python run_gd3.py; tag soict-submission-v4.>')
h2(doc, 'Phụ lục B. Bảng kết quả chi tiết theo từng seed')
body(doc, '<Từ results/reconciliation_canonical.json + raw JSON.>')
h2(doc, 'Phụ lục C. Kiểm định số liệu')
body(doc, '<Quy trình reconcile_verify.py + final_gate_audit.py.>')

doc.save(OUT_DIR / 'KLTN_skeleton.docx')
print('saved', OUT_DIR / 'KLTN_skeleton.docx')
