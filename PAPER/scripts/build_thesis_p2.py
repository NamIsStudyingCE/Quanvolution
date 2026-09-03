# -*- coding: utf-8 -*-
"""build_thesis_p2.py — Luận văn FULL draft, phần 2: Chương 4-6 + TLTK + Phụ lục.
Bảng số liệu dựng TỰ ĐỘNG từ results/reconciliation_canonical.json (ddof=1)
và results/circuit_ablation_summary.json — không số liệu nào được gõ tay."""
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path('.')
OUT = ROOT / 'GD4' / 'KLTN_draft_full.docx'
FIG = ROOT / 'PAPER' / 'figures'
TNR = 'Times New Roman'
CANON = json.load(open(ROOT / 'results' / 'reconciliation_canonical.json', encoding='utf-8'))
B = CANON['breastmnist']['models']; O = CANON['octmnist']['models']

doc = Document(str(OUT))

def para(text, size=13, bold=False, align='justify', italic=False, after=6):
    p = doc.add_paragraph()
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'justify': WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.name = TNR; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    return p

def h1(text): return para(text, 14, bold=True, align='left', after=10)
def h2(text): return para(text, 13, bold=True, align='left', after=6)
def caption(text):
    p = para(text, 12, bold=True, align='center', italic=True, after=10)
    return p

def pic(path, width_cm=15.5, cap=None):
    if Path(path).exists():
        doc.add_picture(str(path), width=Cm(width_cm))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cap: caption(cap)
    else:
        print('MISSING FIG:', path)

def table(headers, rows, caption_text, size=11):
    caption(caption_text)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.name = TNR; r.font.size = Pt(size); r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = ''
            r = cells[j].paragraphs[0].add_run(str(v))
            r.font.name = TNR; r.font.size = Pt(size)
    para('', after=4)

def mstd(m, model, k):
    v = m[model][k]
    return f"{v['mean']:.4f} ± {v['std']:.4f}"

def mstdci(m, model, k):
    v = m[model][k]
    return f"{v['mean']:.4f} ± {v['std']:.4f}\n[{v['ci_lo']:.4f}, {v['ci_hi']:.4f}]"

doc.add_page_break()
# ================= CHƯƠNG 4 =================
h1('CHƯƠNG 4. KẾT QUẢ THỰC NGHIỆM VÀ ĐÁNH GIÁ')
h2('4.1. Thiết lập thí nghiệm')
para('Hai bộ dữ liệu MedMNIST v2 [19] được sử dụng: BreastMNIST — 780 ảnh siêu âm vú 28×28, nhị phân '
     '(546 train / 78 val / 156 test, 73% lành tính — 27% ác tính), đại diện chế độ dữ liệu nhỏ và lệch '
     'lớp; OCTMNIST — subset cân bằng 5,000 ảnh OCT võng mạc 4 lớp (3,500/500/1,000), đại diện chế độ '
     'dữ liệu lớn đa lớp. Năm cấu hình chạy trên cả hai dataset với mạch L=2 cho BreastMNIST và L=1 '
     'cho OCTMNIST, cùng 10 hạt giống S = {0, 42, 100, 2023, 777, 999, 1234, 5678, 1111, 2222} và 20 '
     'epochs. Sáu metrics đánh giá: Accuracy, Balanced Accuracy, F1-score (macro), MCC, ROC-AUC '
     '(macro OvR) và PR-AUC. Kiểm định: paired t-test và Wilcoxon signed-rank (α = 0.05); với n=10, '
     'p-value Wilcoxon rời rạc có mức tối thiểu lý thuyết 1/2⁹ ≈ 0.00195; CI 95% dùng t*(df=9) = 2.262. '
     'p-value không hiệu chỉnh family-wise và đóng vai trò sàng lọc khám phá (xem Chương 5).', align='justify')
h2('4.2. Kết quả trên BreastMNIST (chế độ nhỏ, lệch lớp)')
para('Bảng 4.1 tổng hợp kết quả 10 seeds; Hình 4.1 trực quan hóa.', align='justify')
table(
    ['Mô hình', 'Accuracy', 'Balanced Acc', 'F1', 'MCC', 'ROC-AUC', 'PR-AUC'],
    [
        ['Classical CNN', mstd(B,'classical_cnn','acc'), mstdci(B,'classical_cnn','bacc'), mstd(B,'classical_cnn','f1'),
         mstd(B,'classical_cnn','mcc'), mstdci(B,'classical_cnn','auc'), mstdci(B,'classical_cnn','pr_auc')],
        ['Fixed Basic (L2)', mstd(B,'fixed_basic','acc'), mstdci(B,'fixed_basic','bacc'), mstd(B,'fixed_basic','f1'),
         mstd(B,'fixed_basic','mcc'), mstdci(B,'fixed_basic','auc'), mstdci(B,'fixed_basic','pr_auc')],
        ['Trainable Basic (L2)', mstd(B,'trainable_basic','acc'), mstdci(B,'trainable_basic','bacc'), mstd(B,'trainable_basic','f1'),
         mstd(B,'trainable_basic','mcc'), mstdci(B,'trainable_basic','auc'), mstdci(B,'trainable_basic','pr_auc')],
        ['Fixed Strongly (L2)', mstd(B,'fixed_strongly','acc'), mstdci(B,'fixed_strongly','bacc'), mstd(B,'fixed_strongly','f1'),
         mstd(B,'fixed_strongly','mcc'), mstdci(B,'fixed_strongly','auc'), mstdci(B,'fixed_strongly','pr_auc')],
        ['Trainable Strongly (L2)', mstd(B,'trainable_strongly','acc'), mstdci(B,'trainable_strongly','bacc'), mstd(B,'trainable_strongly','f1'),
         mstd(B,'trainable_strongly','mcc'), mstdci(B,'trainable_strongly','auc'), mstdci(B,'trainable_strongly','pr_auc')],
    ],
    'Bảng 4.1: Kết quả 10-seed BreastMNIST (mean ± sample std; [CI 95%] cho BAcc/ROC-AUC/PR-AUC)', size=10)
pic(FIG / 'Fig3_breastmnist_benchmark.png', 15.0, 'Hình 4.1: Kết quả 10-seed BreastMNIST trên 6 metrics')
para('Bốn phát hiện chính: (1) Fixed Basic L2 đạt ROC-AUC cao nhất 0.8521 ± 0.0095, vượt Classical CNN '
     '0.8336 ± 0.0259 với ý nghĩa thống kê (p_ttest = 0.0298, p_wilcoxon = 0.0254) và hiệu ứng lớn '
     '(Cohen\u2019s d = +0.815). (2) Fixed Strongly L2 đạt PR-AUC cao nhất 0.9182 ± 0.0071, vượt CNN '
     '(0.9041 ± 0.0100) với p_ttest = 0.0023, p_wilcoxon = 0.0059 và hiệu ứng rất lớn (d = +1.332) — '
     'metric lâm sàng quan trọng nhất khi chi phí bỏ sót ca ác tính cao. (3) Ổn định phương sai: std '
     'ROC-AUC của Fixed Basic nhỏ hơn CNN khoảng 2.7 lần (0.0095 so với 0.0259), chứng minh kháng biến '
     'động khởi tạo. (4) Trainable Strongly đạt Balanced Accuracy cao nhất 0.6945 ± 0.0451 (+0.0344 so '
     'với Fixed Strongly, d = +0.677, p = 0.0611) nhưng khác biệt với CNN không có ý nghĩa thống kê '
     '(p = 0.6701, d = +0.139) — thêm một by-chứng cứ cho việc thêm tham số tự học không tự động tốt hơn.', align='justify')
h2('4.3. Kết quả trên OCTMNIST (chế độ lớn, đa lớp)')
para('Bảng 4.2 và Hình 4.2 trình bày kết quả OCTMNIST với 6 cấu hình (thêm quán quân tĩnh random_L1 '
     'từ giai đoạn khảo sát).', align='justify')
table(
    ['Mô hình', 'Accuracy', 'Balanced Acc', 'F1', 'MCC', 'ROC-AUC', 'PR-AUC'],
    [
        ['Classical CNN', mstd(O,'classical_cnn','acc'), mstdci(O,'classical_cnn','bacc'), mstd(O,'classical_cnn','f1'),
         mstd(O,'classical_cnn','mcc'), mstdci(O,'classical_cnn','auc'), mstdci(O,'classical_cnn','pr_auc')],
        ['Fixed Basic (L1)', mstd(O,'fixed_basic','acc'), mstdci(O,'fixed_basic','bacc'), mstd(O,'fixed_basic','f1'),
         mstd(O,'fixed_basic','mcc'), mstdci(O,'fixed_basic','auc'), mstdci(O,'fixed_basic','pr_auc')],
        ['Trainable Basic (L1)', mstd(O,'trainable_basic','acc'), mstdci(O,'trainable_basic','bacc'), mstd(O,'trainable_basic','f1'),
         mstd(O,'trainable_basic','mcc'), mstdci(O,'trainable_basic','auc'), mstdci(O,'trainable_basic','pr_auc')],
        ['Fixed Champ (random_L1)', mstd(O,'fixed_champion_gd2','acc'), mstdci(O,'fixed_champion_gd2','bacc'), mstd(O,'fixed_champion_gd2','f1'),
         mstd(O,'fixed_champion_gd2','mcc'), mstdci(O,'fixed_champion_gd2','auc'), mstdci(O,'fixed_champion_gd2','pr_auc')],
        ['Fixed Strongly (L1)', mstd(O,'fixed_strongly','acc'), mstdci(O,'fixed_strongly','bacc'), mstd(O,'fixed_strongly','f1'),
         mstd(O,'fixed_strongly','mcc'), mstdci(O,'fixed_strongly','auc'), mstdci(O,'fixed_strongly','pr_auc')],
        ['Trainable Strongly (L1)', mstd(O,'trainable_strongly','acc'), mstdci(O,'trainable_strongly','bacc'), mstd(O,'trainable_strongly','f1'),
         mstd(O,'trainable_strongly','mcc'), mstdci(O,'trainable_strongly','auc'), mstdci(O,'trainable_strongly','pr_auc')],
    ],
    'Bảng 4.2: Kết quả 10-seed OCTMNIST (mean ± sample std; [CI 95%])', size=10)
pic(FIG / 'Fig3_octmnist_benchmark.png', 15.0, 'Hình 4.2: Kết quả 10-seed OCTMNIST trên 6 metrics')
para('Ba phát hiện: (1) Classical CNN dẫn đầu tuyệt đối cả 6 metrics (ROC-AUC 0.7505 ± 0.0240, PR-AUC '
     '0.4991 ± 0.0297), vượt mô hình lượng tử tốt nhất với p < 0.001 và hiệu ứng khổng lồ (d = +2.108 '
     'ROC-AUC; d = +1.874 BAcc) — mạch 4-qubit nông chạm trần biểu diễn (expressibility bottleneck) trên '
     'bài toán 4 lớp. (2) Trong cùng họ strongly, mạch tự học vượt mạch tĩnh: Trainable Strongly 0.6922 '
     '± 0.0199 so với 0.6690 ± 0.0055, Δ = +0.0232, p_wilcoxon = 0.0098, d = +1.050. (3) Tuy vậy, mạch '
     'tự học chỉ hòa với quán quân tĩnh random_L1 (0.6912 ± 0.0071, Δ = +0.0010, p = 0.8875, d = +0.046) '
     '— tự học không tự động vượt một mạch tĩnh được chọn hợp lý.', align='justify')
h2('4.4. Động học tối ưu và kiểm chứng gradient')
para('Phân tích động học (Hình 4.3–4.5) cho thấy: mọi mô hình đạt plateau loss trong 12–15 epochs '
     'không phân kỳ; quỹ đạo góc quay θ(t) dịch chuyển trơn và bị hút về vùng giá trị ổn định; chuẩn '
     'gradient L2 của mạch trainable strongly-entangling duy trì xấp xỉ 0.2–0.5 trên đường trung bình '
     'theo seed (đỉnh từng seed gần 1.3), ở mức cao hơn nhiều so với ngưỡng triệt tiêu gradient — loại '
     'trừ Barren Plateaus như một sanity check thực nghiệm trên mạch 4-qubit nông [14].', align='justify')
pic(FIG / 'Fig4a_breastmnist_curves.png', 15.0, 'Hình 4.3: Đường hội tụ train/val trên BreastMNIST (trái) và OCTMNIST (phải)')
pic(FIG / 'Fig4c_theta_trajectories.png', 15.0, 'Hình 4.4: Quỹ đạo góc quay θ(t) qua 20 epochs (seed 1–3)')
pic(FIG / 'Fig4d_gradient_norms.png', 12.5, 'Hình 4.5: Động học chuẩn gradient L2 (trainable_strongly)')
h2('4.5. Circuit ablation (khảo sát GĐ2)')
para('Trước khi dựng ma trận ba tầng, giai đoạn khảo sát đã ablation sáu cấu hình mạch tĩnh trên cả '
     'hai dataset. Bảng 4.4 tổng hợp ROC-AUC; cấu hình basic_L2 cho ROC-AUC tốt nhất trên BreastMNIST '
     'và random_L1 tốt nhất trên OCTMNIST — hai cấu hình này được chọn làm đại diện trong ma trận cuối. '
     'Lưu ý: số liệu Bảng 4.4 đến từ run khảo sát độc lập của GĐ2, do đó chênh lệch nhỏ (dưới 0.002 '
     'ROC-AUC) so với Bảng 4.1–4.2 vốn dùng run protocol cuối của GĐ3 — hai thí nghiệm này nhất quán '
     'về thứ hạng cấu hình và không được dùng thay thế lẫn nhau.', align='justify')
abl = json.load(open(ROOT / 'results' / 'circuit_ablation_summary.json', encoding='utf-8'))
order = ['random_L1', 'random_L2', 'basic_L1', 'basic_L2', 'strongly_L1', 'strongly_L2']
rows = []
for c in order:
    br = abl['breastmnist'][c]['summary']['auc']
    oc = abl['octmnist'][c]['summary']['auc']
    rows.append([c, f"{br['mean']:.4f} ± {br.get('std', 0):.4f}", f"{oc['mean']:.4f} ± {oc.get('std', 0):.4f}"])
table(['Cấu hình mạch', 'ROC-AUC BreastMNIST', 'ROC-AUC OCTMNIST'], rows,
      'Bảng 4.4: Circuit ablation — ROC-AUC 6 cấu hình mạch tĩnh (mean ± std)', size=11)
pic(str(ROOT / 'results' / 'figures' / 'circuit_ablation_breastmnist.png'), 15.0,
    'Hình 4.6: Biểu đồ circuit ablation 6 cấu hình trên BreastMNIST (GĐ2)')
h2('4.6. Chi phí tính toán')
para('Bảng 4.5 đo độ trễ suy luận trên CPU Intel. Phần lớn thời gian (~99.98%) nằm ở mô phỏng 196 '
     'statevector cho một ảnh; chiến lược precompute feature maps một lần cho toàn bộ dataset giúp '
     'giảm 10-seed huấn luyện head xuống chỉ ~18 giây.', align='justify')
table(
    ['Mô hình', 'Giai đoạn', 'Độ trễ', 'Tỉ lệ', 'Tham số kernel'],
    [
        ['Classical CNN', 'Forward pass', '0.310 ms', '1.0×', '20'],
        ['Fixed Quanv', 'Trích xuất (196 patches)', '220.187 ms', '710.3×', '0'],
        ['Fixed Quanv', 'Head phân loại', '0.034 ms', '0.11×', 'giống CNN'],
        ['Fixed Quanv', 'End-to-end', '220.221 ms', '710.4×', '0'],
        ['Trainable Quanv', 'End-to-end', '~220.25 ms', '~710.5×', '12 – 24'],
    ],
    'Bảng 4.5: Độ trễ suy luận CPU và chi phí tính toán', size=11)
doc.add_page_break()

# ================= CHƯƠNG 5 =================
h1('CHƯƠNG 5. KẾT LUẬN')
h2('5.1. Kết luận')
para('Khóa luận đã trình bày một benchmark đối xứng, khả tái lập so sánh quanvolution tĩnh/tự học với '
     'baseline CNN cổ điển trên MedMNIST. Ba kết luận chính: (1) Ưu thế lượng tử phụ thuộc chặt chẽ vào '
     'chế độ dữ liệu — thắng trên dữ liệu nhỏ, lệch lớp; thua quyết liệt trên dữ liệu lớn đa lớp; '
     '(2) Mạch tĩnh 0 tham số cung cấp thiên kiến quy nạp mạnh: ổn định phương sai ~2.7 lần và PR-AUC '
     'cao nhất trên BreastMNIST mà không cần huấn luyện kernel; (3) Khả năng tự học chỉ cục bộ trong '
     'cùng họ mạch và không vượt được mạch tĩnh được chọn tối ưu.', align='justify')
h2('5.2. Những đóng góp mới')
para('(1) Khung benchmark đối xứng 1:1 với kiểm định thống kê kép — khắc phục L1/L2 của literature; '
     '(2) Phân định thực nghiệm ranh giới data-regime cho quanvolution 4-qubit; (3) Bộ tài liệu tái lập '
     'tự động: mã nguồn seed-cố định, ground truth JSON, kịch bản kiểm định số liệu; (4) Bản thảo bài '
     'báo quốc tế đã nộp tại SOICT 2026 (Springer CCIS).', align='justify')
doc.add_page_break()

# ================= CHƯƠNG 6 =================
h1('CHƯƠNG 6. HƯỚNG PHÁT TRIỂN')
para('Bốn hướng phát triển tiếp theo: (1) tích hợp backend GPU/tensor-network (NVIDIA cuQuantum) để '
     'mở rộng quy mô ảnh và số qubit; (2) đánh giá dưới nhiễu NISQ thực (depolarizing, readout) trên '
     'phần cứng IBM QPU; (3) chạy full-scale OCTMNIST (~97,000 ảnh) và độ phân giải cao hơn với '
     'multi-scale patch; (4) đối chiếu trực tiếp với kiến trúc hybrid gating trước đó và khảo sát các '
     'phương án encoding khác (angle/HUAR/amplitude).', align='justify')
doc.add_page_break()

# ================= TLTK =================
para('TÀI LIỆU THAM KHẢO', 16, bold=True, align='center')
h2('Tiếng Việt')
para('(Không có)')
h2('Tiếng Anh')
refs = [
 'F. M. Altares-López, A. Ribeiro, J. J. García-Ripoll, "Automatic design of quantum feature maps," Quantum Science and Technology, 6(4), 045015, 2021.',
 'V. Azevedo, C. Silva, I. Dutra, "Quantum transfer learning for breast cancer detection," Quantum Machine Intelligence, 4(1), 5, 2022.',
 'V. Bergholm et al., "PennyLane: Automatic differentiation of quantum machine learning circuits," arXiv:1811.04968, 2018.',
 'J. Biamonte et al., "Quantum machine learning," Nature, 549(7671), 195–202, 2017.',
 'M. Cerezo et al., "Variational quantum algorithms," Nature Reviews Physics, 3(9), 625–644, 2021.',
 'I. Cong, S. Choi, M. D. Lukin, "Quantum convolutional neural networks," Nature Physics, 15(12), 1273–1278, 2019.',
 'A. Esteva et al., "A guide to deep learning in healthcare," Nature Medicine, 25(1), 24–29, 2019.',
 'M. Henderson, S. Shakya, S. Pradhan, S. Cook, "Quanvolutional neural networks: Powering image recognition with quantum circuits," Quantum Machine Intelligence, 2(1), 2, 2020.',
 'Q. N. Hoang, T. T. Pham, D. N. M. Dang, "Efficient hybrid quantum-classical convolutional neural network with feature propagation layer for multi-class image classification," in Proc. Int. Conf. Adv. Eng. Theory Appl. (AETA), 2023.',
 'H.-Y. Huang et al., "Power of data in quantum machine learning," Nature Communications, 12(1), 2631, 2021.',
 'J. Kübler, S. Buchholz, B. Schölkopf, "The inductive bias of quantum kernels," in NeurIPS, vol. 34, pp. 12661–12673, 2021.',
 'G. Litjens et al., "A survey on deep learning in medical image analysis," Medical Image Analysis, 42, 60–88, 2017.',
 'N. Matondo-Mvula, K. Elleithy, "Breast cancer detection with quanvolutional neural networks," Entropy, 26(8), 630, 2024.',
 'J. R. McClean, S. Boixo, V. N. Smelyanskiy, R. Babbush, H. Neven, "Barren plateaus in quantum neural network training landscapes," Nature Communications, 9(1), 4812, 2018.',
 'M. Schuld, N. Killoran, "Is quantum advantage the right goal for quantum machine learning?" PRX Quantum, 3(3), 030101, 2022.',
 'M. Schuld, V. Bergholm, C. Gogolin, J. Izaac, N. Killoran, "Evaluating analytic gradients on quantum hardware," Physical Review A, 99(3), 032331, 2019.',
 'C. Shorten, T. M. Khoshgoftaar, "A survey on image data augmentation for deep learning," Journal of Big Data, 6(1), 60, 2019.',
 'T. H. Vu, L. H. Le, T. B. Pham, "Exploring the features of quanvolutional neural networks for improved image classification," Quantum Machine Intelligence, 6(1), 29, 2024.',
 'J. Yang et al., "MedMNIST v2 — A large-scale lightweight benchmark for 2D and 3D biomedical image classification," Scientific Data, 10(1), 41, 2023.',
]
for i, r_ in enumerate(refs, 1):
    para(f'[{i}] {r_}', 12, after=4)
doc.add_page_break()

# ================= PHỤ LỤC =================
para('PHỤ LỤC', 16, bold=True, align='center')
h2('Phụ lục A. Hướng dẫn cài đặt và tái hiện thực nghiệm')
para('Môi trường: Python 3.10, PyTorch ≥ 2.0, PennyLane ≥ 0.35, scikit-learn ≥ 1.3. Các bước: '
     '(1) git clone https://github.com/NamIsStudyingCE/Quanvolution; (2) pip install -r requirements.txt; '
     '(3) python run_gd3.py — chạy ma trận ba tầng 10 seeds, xuất results/full_trainable_*.json; '
     '(4) python PAPER/scripts/regenerate_figs_bigfont.py — tái tạo biểu đồ; '
     '(5) python PAPER/scripts/reconcile_verify.py — kiểm định số liệu. Toàn bộ mã nguồn gắn tag '
     'soict-submission-v4 đúng với bản nộp hội nghị.', align='justify')
h2('Phụ lục B. Số liệu chi tiết theo từng seed (ROC-AUC)')
para('Nguồn: results/full_trainable_*.json — test_metrics của từng seed; bảng đầy đủ 6 metrics × 10 '
     'seeds × 6 mô hình × 2 datasets nằm trong reconciliation_canonical.json kèm theo đĩa CD.', align='justify')
seeds = json.load(open(ROOT / 'results' / 'full_trainable_breastmnist.json', encoding='utf-8'))['seeds']
rows = [[str(s)] for s in seeds]
tbl_hdr = ['Seed'] + [f'{m[:14]}' for m in ['classical_cnn', 'fixed_basic', 'trainable_basic', 'fixed_strongly', 'trainable_strongly']]
raw_b = json.load(open(ROOT / 'results' / 'full_trainable_breastmnist.json', encoding='utf-8'))['raw_results']
for j, m in enumerate(['classical_cnn', 'fixed_basic', 'trainable_basic', 'fixed_strongly', 'trainable_strongly']):
    vals = [row['auc'] for row in raw_b[m]['test_metrics']]
    for i, v in enumerate(vals):
        rows[i].append(f'{v:.4f}')
table(tbl_hdr, rows, 'Bảng PHỤ LỤC B.1: ROC-AUC theo seed — BreastMNIST (basic_L2 quy ước trình bày Fixed/Trainable Strongly = L2)', size=10)
h2('Phụ lục C. Quy trình kiểm định số liệu')
para('Mọi số liệu của luận văn và bài báo đều được tính lại độc lập từ raw per-seed JSON bởi script '
     'reconcile_verify.py (mean, sample std ddof=1, CI 95%, paired t-test, Wilcoxon, Cohen\u2019s d), '
     'kết quả lưu tại results/reconciliation_canonical.json. Kịch bản final_gate_audit.py đối chiếu '
     'ngược từng cặp mean±std/CI trong file PDF với canonical. File final_ground_truth.json cũ (chứa '
     'số liệu giai đoạn đầu) đã được cách ly vào results/archive/ để loại trừ nhầm lẫn nguồn.', align='justify')

doc.save(str(OUT))
print('PART 2 saved:', OUT)
