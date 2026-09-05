# -*- coding: utf-8 -*-
"""build_thesis_p3.py — PASS 2: nở nội dung luận văn lên ≥50 trang.
Chèn đoạn văn mới SAU các anchor paragraph (tìm bằng text) — không đụng bảng/số cũ.
Mọi số nêu trong đoạn mới đều trích nguyên văn từ reconciliation_canonical.json."""
import copy
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC = Path('GD4/KLTN_draft_full.docx')
doc = Document(str(DOC))
N_ADD = 0

def find_para(snippet):
    for p in doc.paragraphs:
        if snippet in p.text:
            return p
    raise KeyError(f'anchor not found: {snippet[:60]}')

def ins_after(anchor_p, text, bold=False, align='justify', size=13, heading=False):
    """Create a new paragraph (appended at end) then MOVE its XML right after anchor."""
    global N_ADD
    if isinstance(bold, dict):  # chấp nhận kwargs-dict truyền theo vị trí
        kw = bold
        bold = kw.get('bold', False)
        align = kw.get('align', 'justify')
    p = para(text, size, bold, align, heading)
    anchor_p._p.addnext(p._p)
    N_ADD += 1
    return p

def ins_chain(anchor_p, items):
    """Insert a list of (text, kwargs) sequentially after anchor, preserving order."""
    cur = anchor_p
    for text, kw in items:
        cur = ins_after(cur, text, **kw)

def para(text, size=13, bold=False, align='justify', heading=False):
    p = doc.add_paragraph()
    p.alignment = {'justify': WD_ALIGN_PARAGRAPH.JUSTIFY, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'left': WD_ALIGN_PARAGRAPH.LEFT}[align]
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    r.bold = bold or heading
    return p

def H2(text):
    return (text, {'bold': True, 'align': 'left'})

# ============ CHƯƠNG 1 ============
a = find_para('Chẩn đoán hỗ trợ bằng máy học')
ins_chain(a, [
    ("Trên thực tế, ba modality chẩn đoán hình ảnh phổ biến nhất của đề tài đều mang ràng buộc riêng. "
     "Ảnh siêu âm (BUS) an toàn, không nhiễm xạ, phù hợp sàng lọc phụ nữ có mô tuyến vú đặc, nhưng chịu "
     "nhiễu đốm âm học (speckle noise) và phụ thuộc góc quét của kỹ thuật viên. Ảnh OCT võng mạc cho độ "
     "phân giải cấu trúc các tầng võng mạc cao nhưng mật độ thông tin bệnh lý nằm tập trung ở vùng vàng mạc "
     "trung tâm, đòi hỏi mô hình nắm bắt khác biệt vi thể giữa các giai đoạn bệnh. Cả hai đều là ảnh "
     "grayscale 28×28 trong chuẩn MedMNIST — định dạng nhẹ cho phép benchmark quy mô lớn trên CPU.",
     {'align': 'justify'}),
    ("Vấn đề mất cân bằng lớp trong y học không chỉ là đặc tính dữ liệu mà là hệ quả của dịch tễ học: "
     "ca ác tính luôn là thiểu số so với ca lành tính. Một mô hình chỉ đo bằng Accuracy có thể đạt 90% "
     "khi cứ dự đoán tất cả là lành tính — trong khi bỏ sót 100% ca ác tính, hậu quả lâm sàng thảm khốc. "
     "Vì vậy, bộ metrics của đề tài bắt buộc gồm các thước đo nhạy với lớp thiểu số (Balanced Accuracy, "
     "F1, PR-AUC) và các thước đo chất lượng xếp hạng (ROC-AUC, MCC).", {'align': 'justify'}),
])
a = find_para('Đối tượng: lớp tích chập lượng tử 4-qubit')
ins_after(a, "Bảng tóm tắt phạm vi: hai dataset (BreastMNIST — nhị phân lệch lớp; OCTMNIST — đa lớp "
             "4 nhóm cân bằng); năm cấu hình mô hình đối chứng (CNN, Fixed Basic/Strongly, Trainable "
             "Basic/Strongly) kèm quán quân tĩnh random_L1; protocol 10 seeds × 20 epochs × 6 metrics; "
             "hạ tầng mô phỏng CPU statevector; sản phẩm kèm theo: demo notebook + bản thảo bài báo "
             "quốc tế đã nộp SOICT 2026.")

# ============ CHƯƠNG 2 ============
a = find_para('Học sâu đã trở thành trụ cột của phân tích ảnh y tế')
ins_after(a, "Quy trình điển hình của một hệ CAD gồm bốn bước: thu nhận ảnh (siêu âm, OCT, X-quang), "
             "tiền xử lý (cắt vùng quan tâm, chuẩn hóa), trích xuất đặc trưng, và phân loại/đánh giá. "
             "Trước kỷ nguyên học sâu, bước trích xuất đặc trưng phụ thuộc tay nghề (SIFT, HOG, đặc "
             "trưng hình thái do bác sĩ định nghĩa); CNN đã tự động hóa bước này bằng học biểu diễn "
             "phân tầng. Tuy nhiên, mọi lớp conv bổ sung đều cộng tham số học — món nợ dữ liệu mà các "
             "bộ dữ liệu y tế nhỏ khó trả.", {'align': 'justify'})
a = find_para('QML nghiên cứu việc dùng mạch lượng tử tham số hóa')
ins_chain(a, [
    ("Kiến trúc VQC tiêu chuẩn gồm ba khối: mã hóa (encoding) — nạp dữ liệu cổ điển vào trạng thái "
     "lượng tử bằng các cổng quay tỉ lệ giá trị đặc trưng; lớp biến phân (variational) — chuỗi cổng "
     "quay tham số và cổng rối lặp L tầng; phép đo — kỳ vọng của toán tử Pauli tạo đầu ra cổ điển. "
     "Việc chọn chiến lược mã hóa (angle, amplitude, IQP) và cấu trúc ansatz quyết định khả năng biểu "
     "điễn và độ khó tối ưu.", {'align': 'justify'}),
    ("Hiện tượng Barren Plateaus [14] — gradient trung bình triệt tiêu theo cấp số nhân khi độ sâu và "
     "số qubit tăng — là rủi ro tối ưu hóa lớn nhất của VQC. Với mạch nông 4-qubit của đề tài, lý "
     "thuyết dự báo vùng gradient còn khỏe; xác chứng thực nghiệm của điều này được trình bày tại "
     "Chương 4 như một sanity check bổ trợ.", {'align': 'justify'}),
])
a = find_para('Henderson et al. [8] giới thiệu Quanvolution')
ins_after(a, "Trong thí nghiệm gốc trên MNIST, quanvolution tĩnh (mạch ngẫu nhiên không huấn luyện) "
             "đã đạt độ chính xác 99.96% ngang best-result cổ điển, với hai lợi thế bổ sung: khả năng "
             "huấn luyện tập trung vào head, và tiềm năng khai thác tương quan không gian phi cục bộ "
             "qua rối lượng tử. Tuy nhiên, nghiên cứu gốc chỉ dùng một seed, dataset tổng quát "
             "(không y tế), và không có baseline đối xứng — ba điểm mà đề tài này khắc phục.", {'align': 'justify'})

# ============ CHƯƠNG 3 ============
a = find_para('Pipeline end-to-end gồm bốn giai đoạn tuần tự')
ins_after(a, "Quyết định thiết kế quan trọng nhất của pipeline là \"tầng trích xuất thay thế được, "
             "tầng phân loại bất biến\": cả ba nhóm mô hình (CNN, fixed quanv, trainable quanv) xuất "
             "ra cùng tensor 4×14×14, cùng đi qua một head BatchNorm2d(4)→ReLU→Linear(784→K) được "
             "khởi tạo và huấn luyện theo protocol giống hệt nhau. Mọi chênh lệch hiệu năng giữa các "
             "mô hình do đó chỉ có thể quy về bản chất của phép biến đổi đặc trưng.", {'align': 'justify'})
a = find_para('Protocol huấn luyện chung: CrossEntropy loss')
ins_chain(a, [
    H2('3.7. Bộ chỉ số đánh giá phù hợp dữ liệu y tế'),
    ("Sáu metrics được chọn theo nguyên tắc: mỗi metric trả lời một câu hỏi lâm sàng khác nhau, và "
     "không metric nào đủ đơn lẻ cho dữ liệu lệch lớp. Với nhãn nhị phân, ký hiệu TP/FP/TN/FN theo "
     "quy ước lớp dương = ác tính:", {'align': 'justify'}),
    ("• Accuracy = (TP+TN)/Tổng — tỷ lệ dự đoán đúng toàn phần; tham khảo tổng thể, bị chi phối bởi "
     "lớp đa số.", {'align': 'left'}),
    ("• Balanced Accuracy = (TPR + TNR)/2 — trung bình độ nhạy và độ đặc hiệu; công bằng với lớp "
     "thiểu số.", {'align': 'left'}),
    ("• F1-score (macro) = 2·P·R/(P+R) tính trung bình harmonic trên từng lớp rồi lấy mean — cân bằng "
     "độ chính xác và độ bao phủ.", {'align': 'left'}),
    ("• MCC = (TP·TN − FP·FN)/√((TP+FP)(TP+FN)(TN+FP)(TN+FN)) — tương quan toàn bộ ma trận nhầm lẫn, "
     "khỏi tác động lệch lớp, giá trị [−1,1].", {'align': 'left'}),
    ("• ROC-AUC (OvR macro) — xác suất mô hình xếp một mẫu dương ngẫu nhiên cao hơn một mẫu âm ngẫu "
     "nhiên; độc lập với ngưỡng cắt.", {'align': 'left'}),
    ("• PR-AUC — diện tích dưới đường Precision-Recall; nhạy với lớp dương hiếm, là thước đo lâm sàng "
     "quan trọng nhất khi chi phí bỏ sót ca bệnh cao (đúng đặc thù sàng lọc ung thư).", {'align': 'left'}),
    ("Với OCTMNIST đa lớp, mọi metric nhị phân được mở rộng theo trung bình macro trên các cặp "
     "one-vs-rest; ROC-AUC tính theo One-vs-Rest macro như mô tả tại mục 4.1.", {'align': 'justify'}),
])
a = find_para('Để bảo đảm tính đúng đắn vật lý của gradient analytic statevector')
ins_after(a, "Quy trình kiểm chứng được thực hiện trên cả hai loại mạch (Strongly-Entangling và "
             "Basic-Entangling) với patch ảnh thực từ BreastMNIST/OCTMNIST: gradient analytic "
             "(backprop statevector) và gradient parameter-shift được tính song song trên cùng ma trận "
             "Jacobian; sai lệch tối đa ghi nhận ở cấp 10⁻⁸ — là sai số làm tròn số học, xác nhận hai "
             "đường tính tương đương.", {'align': 'justify'})

# ============ CHƯƠNG 4 ============
a = find_para('Bốn phát hiện chính: (1) Fixed Basic L2 đạt ROC-AUC cao nhất')
ins_chain(a, [
    ("Diễn giải theo từng metric: về Accuracy, mọi mô hình nằm trong biên 0.78–0.81 — sai khác nhỏ, "
     "không đủ phân loại; về Balanced Accuracy, Trainable Strongly dẫn đầu (0.6945) cho thấy mạch tự "
     "học giúp cân bằng độ nhạy hai lớp tốt hơn; về F1 và MCC, Trainable Strongly cũng dẫn đầu "
     "(0.8724, 0.4549) nhưng với phương sai lớn nhất (0.0193, 0.0945) — hiệu năng cao đi kèm bất định; "
     "về hai metric xếp hạng, các mạch tĩnh dẫn đầu với phương sai nhỏ nhất. Ảnh hưởng nổi bật nhất "
     "thuộc về hai mạch tĩnh ở các metric lâm sàng (ROC-AUC, PR-AUC) đúng như định hướng inductive bias.",
     {'align': 'justify'}),
])
a = find_para('Ba phát hiện: (1) Classical CNN dẫn đầu tuyệt đối cả 6 metrics')
ins_after(a, "Diễn giải theo từng metric: Accuracy cả sáu mô hình chụm trong biên 0.39–0.45 (dataset "
             "4 lớp, ngẫu nhiên = 0.25); ROC-AUC và PR-AUC là hai metric tách bạch khoảng cách — CNN "
             "0.7505/0.4991 so với dải lượng tử 0.67–0.69/0.41–0.44; MCC của CNN (0.3156) gần gấp đôi "
             "các mô hình lượng tử (0.2394–0.2566), cho thấy CNN duy trì được tương quan nhãn-đặc trưng "
             "mà mạch nông mất dần. Balanced Accuracy bằng đúng Accuracy trên mọi mô hình — hệ quả của "
             "dataset 4 lớp cân bằng (250 mẫu/lớp test).", {'align': 'justify'})
a = find_para('Diễn giải kết quả ablation: trên BreastMNIST')
ins_after(a, "Đọc kết hợp hai bảng: Bảng 4.1 (L=2 cho BreastMNIST) và Bảng 4.2 (L=1 cho OCTMNIST) — "
             "cùng một head, cùng protocol, hai chế độ dữ liệu đối lập. Bốn cột metric của BreastMNIST "
             "và OCTMNIST được đánh giá độc lập; không có cặp số nào được lấy trung bình chéo dataset. "
             "Kết luận ablation vì vậy mang tính cấu hình (nên dùng mạch nào cho loại dữ liệu nào) "
             "thay vì khái quát hóa tuyệt đối.", {'align': 'justify'})

# ============ CHƯƠNG 5 ============
a = find_para('Khóa luận đã trình bày một benchmark đối xứng, khả tái lập so sánh')
ins_after(a, "Ngoài ba kết luận chính, đề tài còn để lại một tài sản phương pháp luận: quy trình kiểm "
             "định số liệu tự động (reconcile_verify.py, final_gate_audit.py) — mọi bảng số trong luận "
             "văn và bài báo đều được sinh/đối chiếu máy từ raw per-seed JSON, loại trừ triệt để lỗi "
             "sao chép tay. Quy trình này là điều kiện để kết quả nghiên cứu được coi là khả tái lập "
             "theo đúng nguyên tắc số 1 của khóa luận.", align='justify')

# ============ PHỤ LỤC A ============
a = find_para('Môi trường: Python 3.10, PyTorch')
ins_after(a, "Chi tiết cấu trúc đĩa nộp bài (theo quy định Phụ lục 2): thư mục Source chứa mã nguồn "
             "và dữ liệu kết quả; thư mục Document chứa bản mềm luận văn; file README liệt kê corresponding "
             "author và email liên hệ.", align='justify')

doc.save(str(DOC))
print(f'PASS 2 done: {N_ADD} đoạn mới được chèn')
